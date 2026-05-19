#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF/Word/Markdown/JSON 多格式转换脚本
被C++ Qt程序调用，执行实际的转换工作
"""

import sys
import json
import os
import re

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

from pdf2docx import Converter
from docx import Document
import fitz


def _get_cjk_font():
    candidates = [
        "c:/windows/fonts/msyh.ttc",
        "c:/windows/fonts/simsun.ttc",
        "c:/windows/fonts/simhei.ttf",
    ]
    for p in candidates:
        if os.path.exists(p):
            return fitz.Font(fontfile=p)
    return fitz.Font("helv")


class PdfWriter:
    def __init__(self, pdf):
        self.pdf = pdf
        self.font = _get_cjk_font()
        self.page = pdf.new_page()
        self.tw = fitz.TextWriter(self.page.rect)
        self.y = 50
        self.mx = 50
        self.mr = self.page.rect.width - 50

    def _new_page(self):
        self.tw.write_text(self.page)
        self.page = self.pdf.new_page()
        self.tw = fitz.TextWriter(self.page.rect)
        self.y = 50

    def write(self, text, fs=11, indent=0, dy=None):
        if dy is None:
            dy = int(fs * 1.5)
        x = self.mx + indent
        remaining = text
        while remaining:
            if self.y > 780:
                self._new_page()
            avail = self.mr - x
            w = self.font.text_length(remaining, fontsize=fs)
            if w <= avail:
                self.tw.append((x, self.y), remaining, font=self.font, fontsize=fs)
                break
            cut = len(remaining)
            while cut > 1 and self.font.text_length(remaining[:cut], fontsize=fs) > avail:
                cut -= 1
            self.tw.append((x, self.y), remaining[:cut], font=self.font, fontsize=fs)
            remaining = remaining[cut:]
            self.y += dy
            x = self.mx + indent
        self.y += dy

    def nl(self, dy=6):
        self.y += dy
        if self.y > 780:
            self._new_page()

    def finish(self):
        self.tw.write_text(self.page)


def _ensure_space(tw, page, y, needed, margin_x, font, fontsize, max_y=780):
    if y + needed > max_y:
        tw.write_text(page)
        page = pdf_page_new(page.parent)
        tw = _new_writer(page)
        y = 50
    return tw, page, y


def pdf_page_new(doc):
    return doc.new_page()


def convert_pdf_to_word(pdf_path, output_path):
    try:
        if not os.path.exists(pdf_path):
            return {"success": False, "error": f"输入文件不存在: {pdf_path}"}
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        cv = Converter(pdf_path)
        cv.convert(output_path)
        cv.close()
        if os.path.exists(output_path):
            return {"success": True, "output_path": output_path, "message": "转换成功"}
        else:
            return {"success": False, "error": "转换失败，未生成输出文件"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def convert_pdf_to_markdown(pdf_path, output_path):
    try:
        if not os.path.exists(pdf_path):
            return {"success": False, "error": f"输入文件不存在: {pdf_path}"}
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        doc = fitz.open(pdf_path)
        md_lines = []
        for i, page in enumerate(doc):
            blocks = page.get_text("dict")["blocks"]
            md_lines.append(f"## 第 {i+1} 页\n")
            for block in blocks:
                if block["type"] == 0:
                    for line in block["lines"]:
                        text = "".join(span["text"] for span in line["spans"])
                        if not text.strip():
                            continue
                        max_size = max(span["size"] for span in line["spans"])
                        is_bold = any("bold" in span["font"].lower() for span in line["spans"])
                        if max_size >= 18:
                            md_lines.append(f"### {text.strip()}\n")
                        elif is_bold:
                            md_lines.append(f"**{text.strip()}**\n")
                        else:
                            md_lines.append(f"{text.strip()}\n")
            md_lines.append("")
        doc.close()
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(md_lines))
        if os.path.exists(output_path):
            return {"success": True, "output_path": output_path, "message": "转换成功"}
        else:
            return {"success": False, "error": "转换失败，未生成输出文件"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def convert_word_to_pdf(word_path, output_path):
    try:
        if not os.path.exists(word_path):
            return {"success": False, "error": f"输入文件不存在: {word_path}"}
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        try:
            from docx2pdf import convert
            convert(word_path, output_path)
            if os.path.exists(output_path):
                return {"success": True, "output_path": output_path, "message": "转换成功"}
        except Exception:
            pass
        doc = Document(word_path)
        pdf = fitz.open()
        w = PdfWriter(pdf)
        for para in doc.paragraphs:
            text = para.text
            if not text.strip():
                w.nl(10)
                continue
            is_heading = para.style.name.startswith("Heading")
            fs = 16 if is_heading else 11
            w.write(text, fs=fs)
        w.finish()
        pdf.save(output_path)
        pdf.close()
        if os.path.exists(output_path):
            return {"success": True, "output_path": output_path, "message": "转换成功(文本模式)"}
        else:
            return {"success": False, "error": "转换失败，未生成输出文件"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def convert_markdown_to_pdf(md_path, output_path):
    try:
        if not os.path.exists(md_path):
            return {"success": False, "error": f"输入文件不存在: {md_path}"}
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        with open(md_path, "r", encoding="utf-8") as f:
            lines = f.readlines()

        pdf = fitz.open()
        w = PdfWriter(pdf)

        for line in lines:
            raw = line.rstrip("\n")
            stripped = raw.strip()

            if not stripped:
                w.nl(8)
                continue
            if stripped.startswith("```"):
                continue

            if stripped.startswith("# "):
                w.write(stripped[2:], fs=18, dy=28)
                continue
            if stripped.startswith("## "):
                w.write(stripped[3:], fs=15, dy=24)
                continue
            if stripped.startswith("### "):
                w.write(stripped[4:], fs=13, dy=20)
                continue

            if stripped.startswith("> "):
                w.write(stripped[2:], fs=11, indent=20, dy=17)
                continue

            if stripped.startswith("- ") or stripped.startswith("* "):
                w.write("•  " + stripped[2:], fs=11, indent=10, dy=17)
                continue

            m = re.match(r'^(\d+)\.\s+(.*)', stripped)
            if m:
                w.write(f"{m.group(1)}.  {m.group(2)}", fs=11, indent=10, dy=17)
                continue

            if stripped.startswith("|") and stripped.endswith("|"):
                cells = [c.strip() for c in stripped.split("|")[1:-1]]
                if all(set(c) <= set("-: ") for c in cells):
                    continue
                w.write("    ".join(cells), fs=10, indent=5, dy=15)
                continue

            w.write(stripped, fs=11, dy=17)

        w.finish()
        pdf.save(output_path)
        pdf.close()
        if os.path.exists(output_path):
            return {"success": True, "output_path": output_path, "message": "转换成功"}
        else:
            return {"success": False, "error": "转换失败，未生成输出文件"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def convert_pdf_to_json(pdf_path, output_path):
    try:
        if not os.path.exists(pdf_path):
            return {"success": False, "error": f"输入文件不存在: {pdf_path}"}
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        doc = fitz.open(pdf_path)
        result = {"total_pages": len(doc), "file_size": os.path.getsize(pdf_path), "pages": []}
        for i, page in enumerate(doc):
            blocks = page.get_text("dict")["blocks"]
            page_data = {"page_number": i + 1, "width": page.rect.width, "height": page.rect.height, "blocks": []}
            for block in blocks:
                if block["type"] == 0:
                    block_data = {"type": "text", "lines": []}
                    for line in block["lines"]:
                        line_data = {"spans": []}
                        for span in line["spans"]:
                            line_data["spans"].append({
                                "text": span["text"],
                                "font": span["font"],
                                "size": round(span["size"], 2),
                                "color": f"#{span['color']:06x}",
                                "bbox": [round(v, 2) for v in span["bbox"]]
                            })
                        block_data["lines"].append(line_data)
                    page_data["blocks"].append(block_data)
                elif block["type"] == 1:
                    page_data["blocks"].append({"type": "image", "bbox": [round(v, 2) for v in block["bbox"]]})
            result["pages"].append(page_data)
        doc.close()
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        if os.path.exists(output_path):
            return {"success": True, "output_path": output_path, "message": "转换成功"}
        else:
            return {"success": False, "error": "转换失败，未生成输出文件"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def convert_json_to_pdf(json_path, output_path):
    try:
        if not os.path.exists(json_path):
            return {"success": False, "error": f"输入文件不存在: {json_path}"}
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        pdf = fitz.open()
        w = PdfWriter(pdf)
        for line in json.dumps(data, ensure_ascii=False, indent=2).split("\n"):
            w.write(line, fs=10, dy=14)
        w.finish()
        pdf.save(output_path)
        pdf.close()
        if os.path.exists(output_path):
            return {"success": True, "output_path": output_path, "message": "转换成功"}
        else:
            return {"success": False, "error": "转换失败，未生成输出文件"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def convert_word_to_json(word_path, output_path):
    try:
        if not os.path.exists(word_path):
            return {"success": False, "error": f"输入文件不存在: {word_path}"}
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        doc = Document(word_path)
        result = {"paragraphs": [], "tables": []}
        for para in doc.paragraphs:
            if para.text.strip():
                result["paragraphs"].append({
                    "style": para.style.name,
                    "text": para.text,
                    "bold": any(r.bold for r in para.runs if r.bold),
                })
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            result["tables"].append(table_data)
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        if os.path.exists(output_path):
            return {"success": True, "output_path": output_path, "message": "转换成功"}
        else:
            return {"success": False, "error": "转换失败，未生成输出文件"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def convert_json_to_word(json_path, output_path):
    try:
        if not os.path.exists(json_path):
            return {"success": False, "error": f"输入文件不存在: {json_path}"}
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        doc = Document()
        if isinstance(data, dict):
            if "paragraphs" in data:
                for p in data["paragraphs"]:
                    text = p.get("text", "")
                    style = p.get("style", "Normal")
                    para = doc.add_paragraph(text, style=style)
                    if p.get("bold"):
                        for run in para.runs:
                            run.bold = True
            if "tables" in data:
                for table_data in data["tables"]:
                    if table_data:
                        cols = len(table_data[0])
                        table = doc.add_table(rows=len(table_data), cols=cols)
                        for i, row in enumerate(table_data):
                            for j, cell in enumerate(row):
                                table.rows[i].cells[j].text = str(cell)
            if "paragraphs" not in data and "tables" not in data:
                doc.add_paragraph(json.dumps(data, ensure_ascii=False, indent=2))
        else:
            doc.add_paragraph(json.dumps(data, ensure_ascii=False, indent=2))
        doc.save(output_path)
        if os.path.exists(output_path):
            return {"success": True, "output_path": output_path, "message": "转换成功"}
        else:
            return {"success": False, "error": "转换失败，未生成输出文件"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def convert_word_to_markdown(word_path, output_path):
    try:
        if not os.path.exists(word_path):
            return {"success": False, "error": f"输入文件不存在: {word_path}"}
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        doc = Document(word_path)
        md_lines = []
        for para in doc.paragraphs:
            text = para.text
            if not text.strip():
                md_lines.append("")
                continue
            style = para.style.name
            if style.startswith("Heading 1"):
                md_lines.append(f"# {text}\n")
            elif style.startswith("Heading 2"):
                md_lines.append(f"## {text}\n")
            elif style.startswith("Heading 3"):
                md_lines.append(f"### {text}\n")
            elif style.startswith("List"):
                md_lines.append(f"- {text}")
            else:
                parts = []
                for run in para.runs:
                    if run.bold:
                        parts.append(f"**{run.text}**")
                    elif run.italic:
                        parts.append(f"*{run.text}*")
                    else:
                        parts.append(run.text)
                md_lines.append("".join(parts))
        for table in doc.tables:
            md_lines.append("")
            headers = [cell.text for cell in table.rows[0].cells]
            md_lines.append("| " + " | ".join(headers) + " |")
            md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in table.rows[1:]:
                cells = [cell.text for cell in row.cells]
                md_lines.append("| " + " | ".join(cells) + " |")
            md_lines.append("")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(md_lines))
        if os.path.exists(output_path):
            return {"success": True, "output_path": output_path, "message": "转换成功"}
        else:
            return {"success": False, "error": "转换失败，未生成输出文件"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def convert_markdown_to_word(md_path, output_path):
    try:
        if not os.path.exists(md_path):
            return {"success": False, "error": f"输入文件不存在: {md_path}"}
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        with open(md_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
        doc = Document()
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r'^\d+\.\s+', stripped):
                text = re.sub(r'^\d+\.\s+', '', stripped)
                doc.add_paragraph(text, style="List Number")
            elif stripped.startswith("|"):
                pass
            elif stripped.startswith("> "):
                doc.add_paragraph(stripped[2:], style="Quote")
            elif stripped.startswith("```"):
                continue
            else:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', stripped)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith("`") and part.endswith("`"):
                        run = p.add_run(part[1:-1])
                        run.font.name = "Courier New"
                    else:
                        p.add_run(part)
        doc.save(output_path)
        if os.path.exists(output_path):
            return {"success": True, "output_path": output_path, "message": "转换成功"}
        else:
            return {"success": False, "error": "转换失败，未生成输出文件"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def parse_page_ranges(spec, total_pages):
    pages = set()
    for part in spec.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-", 1)
            a = max(1, int(a))
            b = min(total_pages, int(b))
            pages.update(range(a, b + 1))
        else:
            p = int(part)
            if 1 <= p <= total_pages:
                pages.add(p)
    return sorted(pages)


def split_pdf(pdf_path, output_dir, pages_per_file=None, page_ranges=None):
    try:
        if not os.path.exists(pdf_path):
            return {"success": False, "error": f"输入文件不存在: {pdf_path}"}
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        doc = fitz.open(pdf_path)
        total = len(doc)
        base = os.path.splitext(os.path.basename(pdf_path))[0]
        output_files = []
        if page_ranges:
            pages = parse_page_ranges(page_ranges, total)
            out = fitz.open()
            for p in pages:
                out.insert_pdf(doc, from_page=p - 1, to_page=p - 1)
            out_path = os.path.join(output_dir, f"{base}_selected.pdf")
            out.save(out_path)
            out.close()
            output_files.append(out_path)
        elif pages_per_file:
            pages_per_file = int(pages_per_file)
            part = 1
            for start in range(0, total, pages_per_file):
                end = min(start + pages_per_file - 1, total - 1)
                out = fitz.open()
                out.insert_pdf(doc, from_page=start, to_page=end)
                out_path = os.path.join(output_dir, f"{base}_part{part}.pdf")
                out.save(out_path)
                out.close()
                output_files.append(out_path)
                part += 1
        else:
            for i in range(total):
                out = fitz.open()
                out.insert_pdf(doc, from_page=i, to_page=i)
                out_path = os.path.join(output_dir, f"{base}_page{i + 1}.pdf")
                out.save(out_path)
                out.close()
                output_files.append(out_path)
        doc.close()
        return {"success": True, "output_files": output_files, "total_pages": total, "split_count": len(output_files), "message": f"拆分完成，共{len(output_files)}个文件"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def merge_pdfs(pdf_paths, output_path):
    try:
        for p in pdf_paths:
            if not os.path.exists(p):
                return {"success": False, "error": f"文件不存在: {p}"}
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        merged = fitz.open()
        total_pages = 0
        for p in pdf_paths:
            doc = fitz.open(p)
            merged.insert_pdf(doc)
            total_pages += len(doc)
            doc.close()
        merged.save(output_path)
        merged.close()
        if os.path.exists(output_path):
            return {"success": True, "output_path": output_path, "total_pages": total_pages, "file_count": len(pdf_paths), "message": f"合并完成，共{total_pages}页"}
        else:
            return {"success": False, "error": "合并失败，未生成输出文件"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def extract_text_preview(pdf_path, max_pages=3):
    try:
        if not os.path.exists(pdf_path):
            return {"success": False, "error": f"文件不存在: {pdf_path}"}
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        text_content = []
        for page_num in range(min(max_pages, total_pages)):
            page = doc[page_num]
            text = page.get_text()
            if text:
                text_content.append({"page": page_num + 1, "content": text})
        doc.close()
        return {"success": True, "pages": text_content, "total_pages": total_pages}
    except Exception as e:
        return {"success": False, "error": str(e)}


def get_pdf_info(pdf_path):
    try:
        if not os.path.exists(pdf_path):
            return {"success": False, "error": f"文件不存在: {pdf_path}"}
        doc = fitz.open(pdf_path)
        info = {"success": True, "total_pages": len(doc), "file_size": os.path.getsize(pdf_path)}
        doc.close()
        return info
    except Exception as e:
        return {"success": False, "error": str(e)}


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "缺少命令参数"}))
        sys.exit(1)

    command = sys.argv[1]

    try:
        if command == "convert":
            if len(sys.argv) < 4:
                print(json.dumps({"error": "convert命令需要pdf_path和output_path参数"}))
                sys.exit(1)
            result = convert_pdf_to_word(sys.argv[2], sys.argv[3])
            print(json.dumps(result, ensure_ascii=False))

        elif command == "pdf2md":
            if len(sys.argv) < 4:
                print(json.dumps({"error": "pdf2md命令需要pdf_path和output_path参数"}))
                sys.exit(1)
            result = convert_pdf_to_markdown(sys.argv[2], sys.argv[3])
            print(json.dumps(result, ensure_ascii=False))

        elif command == "word2pdf":
            if len(sys.argv) < 4:
                print(json.dumps({"error": "word2pdf命令需要word_path和output_path参数"}))
                sys.exit(1)
            result = convert_word_to_pdf(sys.argv[2], sys.argv[3])
            print(json.dumps(result, ensure_ascii=False))

        elif command == "md2pdf":
            if len(sys.argv) < 4:
                print(json.dumps({"error": "md2pdf命令需要md_path和output_path参数"}))
                sys.exit(1)
            result = convert_markdown_to_pdf(sys.argv[2], sys.argv[3])
            print(json.dumps(result, ensure_ascii=False))

        elif command == "pdf2json":
            if len(sys.argv) < 4:
                print(json.dumps({"error": "pdf2json命令需要pdf_path和output_path参数"}))
                sys.exit(1)
            result = convert_pdf_to_json(sys.argv[2], sys.argv[3])
            print(json.dumps(result, ensure_ascii=False))

        elif command == "json2pdf":
            if len(sys.argv) < 4:
                print(json.dumps({"error": "json2pdf命令需要json_path和output_path参数"}))
                sys.exit(1)
            result = convert_json_to_pdf(sys.argv[2], sys.argv[3])
            print(json.dumps(result, ensure_ascii=False))

        elif command == "word2json":
            if len(sys.argv) < 4:
                print(json.dumps({"error": "word2json命令需要word_path和output_path参数"}))
                sys.exit(1)
            result = convert_word_to_json(sys.argv[2], sys.argv[3])
            print(json.dumps(result, ensure_ascii=False))

        elif command == "json2word":
            if len(sys.argv) < 4:
                print(json.dumps({"error": "json2word命令需要json_path和output_path参数"}))
                sys.exit(1)
            result = convert_json_to_word(sys.argv[2], sys.argv[3])
            print(json.dumps(result, ensure_ascii=False))

        elif command == "word2md":
            if len(sys.argv) < 4:
                print(json.dumps({"error": "word2md命令需要word_path和output_path参数"}))
                sys.exit(1)
            result = convert_word_to_markdown(sys.argv[2], sys.argv[3])
            print(json.dumps(result, ensure_ascii=False))

        elif command == "md2word":
            if len(sys.argv) < 4:
                print(json.dumps({"error": "md2word命令需要md_path和output_path参数"}))
                sys.exit(1)
            result = convert_markdown_to_word(sys.argv[2], sys.argv[3])
            print(json.dumps(result, ensure_ascii=False))

        elif command == "split":
            if len(sys.argv) < 3:
                print(json.dumps({"error": "split命令需要pdf_path参数"}))
                sys.exit(1)
            pdf_path = sys.argv[2]
            output_dir = sys.argv[3] if len(sys.argv) > 3 else os.path.dirname(pdf_path)
            pages_per_file = sys.argv[4] if len(sys.argv) > 4 else None
            page_ranges = sys.argv[5] if len(sys.argv) > 5 else None
            result = split_pdf(pdf_path, output_dir, pages_per_file, page_ranges)
            print(json.dumps(result, ensure_ascii=False))

        elif command == "merge":
            if len(sys.argv) < 4:
                print(json.dumps({"error": "merge命令需要output_path和至少一个pdf_path参数"}))
                sys.exit(1)
            output_path = sys.argv[2]
            pdf_paths = sys.argv[3:]
            result = merge_pdfs(pdf_paths, output_path)
            print(json.dumps(result, ensure_ascii=False))

        elif command == "preview":
            if len(sys.argv) < 3:
                print(json.dumps({"error": "preview命令需要pdf_path参数"}))
                sys.exit(1)
            pdf_path = sys.argv[2]
            max_pages = int(sys.argv[3]) if len(sys.argv) > 3 else 3
            result = extract_text_preview(pdf_path, max_pages)
            print(json.dumps(result, ensure_ascii=False))

        elif command == "info":
            if len(sys.argv) < 3:
                print(json.dumps({"error": "info命令需要pdf_path参数"}))
                sys.exit(1)
            result = get_pdf_info(sys.argv[2])
            print(json.dumps(result, ensure_ascii=False))

        else:
            print(json.dumps({"error": f"未知命令: {command}"}))
            sys.exit(1)

    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)


if __name__ == "__main__":
    main()
